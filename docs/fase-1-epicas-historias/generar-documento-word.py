# Genera "Fase 1 - Epicas e Historias de Usuario - Floristeria.docx" a partir de los .md
# de esta carpeta. El .docx NO se versiona (ver .gitignore); se regenera cuando se
# necesite entregar el documento.
#
# Requisitos: pip install python-docx
# Uso:        python generar-documento-word.py
#
# El script siempre lee el contenido actual de los .md, así que el Word nunca queda
# desactualizado respecto al repositorio: basta con volver a ejecutarlo.

import re
import glob
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
GREEN = RGBColor(0x1F, 0x6B, 0x3A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base styles ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = GREEN

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(hdr_cells[i], "1F6B3A")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table

def parse_md_table(md_text, start_marker=None):
    """Return list of rows (list of cells) from the first markdown table found."""
    lines = md_text.splitlines()
    rows = []
    in_table = False
    for line in lines:
        if line.strip().startswith('|'):
            in_table = True
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if re.match(r'^:?-+:?$', cells[0].replace(' ', '')):
                continue
            rows.append(cells)
        elif in_table:
            break
    return rows

def clean_inline(text):
    text = text.replace('**', '')
    text = text.replace('`', '')
    return text

def add_markdown_body(md_text, skip_h1=True, base_level=2):
    """Render a subset of markdown (headings, tables, bullet lists, paragraphs) into doc."""
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = clean_inline(stripped.lstrip('#').strip())
            if level == 1 and skip_h1:
                i += 1
                continue
            doc.add_heading(text, level=min(level + base_level - 1, 4))
            i += 1
            continue
        if stripped.startswith('>'):
            p = doc.add_paragraph(clean_inline(stripped.lstrip('>').strip()))
            p.style = doc.styles['Intense Quote'] if 'Intense Quote' in [s.name for s in doc.styles] else p.style
            i += 1
            continue
        if stripped.startswith('|'):
            table_lines = []
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_md_table('\n'.join(table_lines))
            if rows:
                add_table(rows[0], rows[1:])
                doc.add_paragraph()
            continue
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(clean_inline(stripped[2:]), style='List Bullet')
            i += 1
            continue
        if re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(clean_inline(re.sub(r'^\d+\.\s', '', stripped)), style='List Number')
            i += 1
            continue
        # plain paragraph (merge consecutive lines)
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not lines[i].strip().startswith(('#', '|', '-', '*', '>')):
            para_lines.append(lines[i].strip())
            i += 1
        doc.add_paragraph(clean_inline(' '.join(para_lines)))

# ================= PORTADA =================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Fase 1 — Épicas e Historias de Usuario')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = GREEN

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Sistema de tienda virtual a la medida para una floristería')
run.font.size = Pt(15)
run.font.color.rgb = GREY

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run(
    'Proyecto formativo SENA — Análisis y Desarrollo de Software\n'
    'Negocio de referencia (datos genéricos): [Floristería Aroma de Rosas]\n'
    'Objetivo futuro: entregar el software a una floristería real en Neiva, Huila\n'
    f'Fecha: {date.today().strftime("%d de %B de %Y")}'
)
info_run.font.size = Pt(11)
info_run.font.color.rgb = GREY

doc.add_page_break()

# ================= TABLA DE CONTENIDO (manual) =================
doc.add_heading('Contenido', level=1)
toc_items = [
    '1. Definición del proyecto',
    '2. Épicas',
    '3. Backlog de historias de usuario',
    '4. Asignación de trabajo del equipo',
    '5. Detalle de historias de usuario (HU-001 a HU-043)',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ================= 1-4: documentos base =================
sections = [
    ('01-definicion-del-proyecto.md', '1. Definición del proyecto'),
    ('02-epicas.md', '2. Épicas'),
    ('03-backlog-historias.md', '3. Backlog de historias de usuario'),
    ('04-asignacion-equipo.md', '4. Asignación de trabajo del equipo'),
]

for filename, heading in sections:
    with open(f'{BASE}/{filename}', encoding='utf-8') as f:
        content = f.read()
    doc.add_heading(heading, level=1)
    add_markdown_body(content, skip_h1=True, base_level=2)
    doc.add_page_break()

# ================= 5: historias de usuario =================
doc.add_heading('5. Detalle de historias de usuario', level=1)
doc.add_paragraph(
    'A continuación, el detalle completo de cada historia de usuario del backlog, en '
    'formato Dado/Cuando/Entonces, con sus reglas de negocio y notas de dependencia.'
)

hu_files = sorted(glob.glob(f'{BASE}/historias-de-usuario/hu-*.md'),
                   key=lambda p: int(re.search(r'hu-(\d+)', p).group(1)))

for path in hu_files:
    with open(path, encoding='utf-8') as f:
        content = f.read()
    lines = content.splitlines()
    h1 = lines[0].lstrip('#').strip()
    doc.add_heading(h1, level=2)
    add_markdown_body('\n'.join(lines[1:]), skip_h1=False, base_level=2)

out_path = f'{BASE}/Fase 1 - Epicas e Historias de Usuario - Floristeria.docx'
doc.save(out_path)
print('OK ->', out_path)
